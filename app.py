import streamlit as st
from docx import Document
from datetime import datetime
import os
import platform
import subprocess
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK


def apply_formatting(run, font_name, font_size, bold=False):
    """Apply specific formatting to a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold

def replace_and_format(doc, placeholders, font_name, font_size, option):
    """Replace placeholders and apply formatting."""
    for para in doc.paragraphs:
        if para.text:  # Check if paragraph has text
            for key, value in placeholders.items():
                if key in para.text:
                    runs = para.runs
                    for run in runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
                            if para == doc.paragraphs[0]:  # Bold replacements only in the first paragraph
                                apply_formatting(run, font_name, font_size, bold=True)
                        else:
                            run.text = run.text.replace(key, value)  # For other paragraphs

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():  # Ensure cell is not empty
                    for key, value in placeholders.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if key == "<<Address>>" else WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    apply_formatting(run, "Times New Roman", 11 if option == "NDA" else 12)  # Ensure consistent formatting
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for para in doc.paragraphs:
        if "Signature Details:" in para.text:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Keep "Signature Details:" left-aligned
            for run in para.runs:
                run.font.size = Pt(11)  # Ensure consistent font size
        elif any(placeholder in para.text for placeholder in placeholders.keys()):
            for key, value in placeholders.items():
                if key in para.text:
                    para.text = para.text.replace(key, value)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align placeholders where necessary
                    for run in para.runs:
                        run.font.size = Pt(11)  # Ensure consistent font size
        if "<< Date >>" in para.text:
            for run in para.runs:
                if "<< Date >>" in run.text:
                    run.text = run.text.replace("<< Date >>", placeholders.get("<< Date >>", ""))
                    apply_formatting(run, "Times New Roman", 12, bold=True)

def edit_word_template(template_path, output_path, placeholders, font_name, font_size, option):
    """Edit Word document and apply formatting."""
    try:
        doc = Document(template_path)

        # Replace placeholders and format
        replace_and_format(doc, placeholders, font_name, font_size, option)

        # Save the modified document
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Error editing Word template: {e}")

def choose_template(currency, include_digital_services):
    """Select the appropriate template based on currency and digital services inclusion."""
    templates = {
        "USD": {
            True: "DM & Automations Services Pricing - USD.docx",
            False: "DM & Automations Services Pricing - USD (without digital service).docx",
        },
        "Rupees": {
            True: "DM & Automations Services Pricing - Rupees.docx",
            False: "DM & Automations Services Pricing - Rupees (without digital service).docx",
        },
        "Pounds": {
            True: "DM & Automations Services Pricing - Pounds.docx",
            False: "DM & Automations Services Pricing - Pounds (without digital service).docx",
        },
    }
    return templates[currency][include_digital_services]    

def edit_pricing_template(template_path, output_path, name, designation, contact, email, location, selected_services):
    try:
        
        digital_marketing_services = [
            "Marketing Strategy",
            "Social Media Channels",
            "Creatives (10 Per Month)",
            "Creatives (20 Per Month)",
            "Creatives (30 Per Month)",
            "Reels (10 Reels)",
            "Meta Ad Account Setup & Pages Setup",
            "Paid Ads (Lead Generation)",
            "Monthly Maintenance & Reporting",
        ]


        # Check if at least one digital marketing service is selected
        include_digital_services = any(service in selected_services for service in digital_marketing_services)

        template_path = choose_template(currency, include_digital_services)
        all_services_selected = set(selected_services) >= set(digital_marketing_services)

        # Load and update the chosen template
        doc = Document(template_path)        

        # Replace placeholders in the general paragraphs
        for para in doc.paragraphs:
            if "<<Client Name>>" in para.text:
                para.text = para.text.replace("<<Client Name>>", name)
            if "<<Client Designation>>" in para.text:
                para.text = para.text.replace("<<Client Designation>>", designation)
            if "<<Client Contact>>" in para.text:
                para.text = para.text.replace("<<Client Contact>>", contact)
            if "<<Client Email>>" in para.text:
                para.text = para.text.replace("<<Client Email>>", email)
            if "<<Client Location>>" in para.text:
                para.text = para.text.replace("<<Client Location>>", location)
            if "<< Date >>" in para.text:
                for run in para.runs:
                    if "<< Date >>" in run.text:
            # Use the input date and format it as "dd/mm/yyyy"
                       formatted_date = date_field.strftime("%d/%m/%Y")
                       run.text = run.text.replace("<< Date >>", formatted_date)
                       apply_formatting(run, "Times New Roman", 12, bold=True)

                        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "<<Client Name>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Name>>", name)
                    if "<<Client Designation>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Designation>>", designation)
                    if "<<Client Contact>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Contact>>", contact)
                    if "<<Client Email>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Email>>", email)
                    if "<<Client Location>>" in cell.text:
                        cell.text = cell.text.replace("<<Client Location>>", location)
        # Process tables to find and update the SPOC table and service table separately
        spoc_table_found = False

        for para in doc.paragraphs:
            if "Supporting SPOC Details" in para.text:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the heading
                spoc_table_found = True

        for table_idx, table in enumerate(doc.tables):
            if spoc_table_found and table_idx == 0:  # Assuming SPOC table is the first table after the identifier
                for row in table.rows:
                    if "Project Sponsor/Client’s Detail" in row.cells[0].text:
                        row.cells[1].text = name
                        row.cells[2].text = designation
                        row.cells[3].text = contact
                        row.cells[4].text = email
                    # Set alignment and font style for cells
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                                run.font.size = Pt(10)  # Set font size smaller
                spoc_table_found = False
            else:
                rows_to_delete = []
                for row_idx, row in enumerate(table.rows[1:], start=1):
                    service_name = row.cells[0].text.strip()
                    if service_name not in selected_services:
                        rows_to_delete.append(row_idx)

                for row_idx in reversed(rows_to_delete):
                    table._element.remove(table.rows[row_idx]._element)
                    
        # Handle "Next Steps" dynamically based on the service table position
        service_table_found = False
        for table in doc.tables:
            if "Name" in table.rows[0].cells[0].text:  # Identify the service table
                service_table_found = True
                table_position = table._element.getparent().index(table._element)
                break

        # Adjust the "Next Steps" section
        next_steps_found = False
        for idx, para in enumerate(doc.paragraphs):
            if "Next Steps:" in para.text:
                next_steps_found = True
                # Check if the service table and "Next Steps" are on the same page
                if service_table_found:
                    # Avoid unnecessary page breaks
                    if idx - table_position > 10:  # Adjust this threshold based on your layout
                        page_break = para.insert_paragraph_before()
                        run = page_break.add_run()
                        run.add_break(WD_BREAK.PAGE)
                break

        # Save the updated document
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Error editing Word template: {e}")
    
def options_changed():
    if "current_input" not in st.session_state:
        return False
    return st.session_state["current_input"] != current_input  
current_input = {}

def format_date_with_suffix(date_obj):
    day = date_obj.day
    month = date_obj.strftime("%B")
    year = date_obj.year
    
    # Determine the suffix for the day
    if 10 <= day % 100 <= 20:  # Special case for 11th, 12th, 13th, etc.
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    
    # Return the formatted date
    return f"{day}{suffix} {month} {year}"

# File to persist invoice number
INVOICE_FILE = "invoice_counter.txt"

if "download_visible" not in st.session_state:
    st.session_state.download_visible = False

if "last_inputs" not in st.session_state:
    st.session_state.last_inputs = {}

if "first_run" not in st.session_state:
    st.session_state.first_run = True

if "output_path" not in st.session_state:
    st.session_state.output_path = None

def track_changes(inputs):
    """Check if the current inputs differ from the last saved inputs."""
    if st.session_state.first_run:
        st.session_state.first_run = False
    else:
        # Compare current inputs with the stored ones after the first run
        if st.session_state.last_inputs != inputs:
            st.session_state.download_visible = False
    # Store the current inputs
    st.session_state.last_inputs = inputs

def get_next_invoice_number():
    """Fetch and increment the invoice number."""
    if not os.path.exists(INVOICE_FILE):
        with open(INVOICE_FILE, "w") as file:
            file.write("1000")  # Starting invoice number

    try:
        with open(INVOICE_FILE, "r") as file:
            content = file.read().strip()
            current_invoice = int(content) if content else 1000
    except ValueError:
        current_invoice = 1000

    next_invoice = current_invoice + 1
    with open(INVOICE_FILE, "w") as file:
        file.write(str(next_invoice))

    return current_invoice

def replace_placeholders(doc, placeholders):
    """Replace placeholders in a document while maintaining proper alignment."""
    # Keywords to detect left-side content
    left_side_keywords = [
        "BILL TO", "Mobile No", "Address", "Email", "Project Name", "Company Name"
    ]

    # Iterate through all paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                inline = para.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        inline[i].text = inline[i].text.replace(key, value)
                # Force left alignment for specific placeholders
                if any(keyword in para.text for keyword in left_side_keywords):
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.left_indent = None  # Reset any indent
                    para.paragraph_format.first_line_indent = None  # Reset first-line indent

    # Iterate through all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in placeholders.items():
                        if key in para.text:
                            inline = para.runs
                            for i in range(len(inline)):
                                if key in inline[i].text:
                                    inline[i].text = inline[i].text.replace(key, value)
                            # Force left alignment for specific placeholders in tables
                            if any(keyword in para.text for keyword in left_side_keywords):
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                para.paragraph_format.left_indent = None
                                para.paragraph_format.first_line_indent = None

    return doc


def format_percentage(value):
    """Format percentage without decimals."""
    return f"{int(value)}%"


def edit_invoice_template(template_name, output_path, placeholders):
    """Edit an invoice template and save the result."""
    try:
        doc = Document(template_name)
        replace_placeholders(doc, placeholders)
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Error editing invoice template: {e}")

def format_price(price, currency):
    """Format price to display correctly with the currency."""
    if price.is_integer():
        formatted_price = f"{int(price)}"
    else:
        formatted_price = f"{price:.2f}"
    if currency == "USD":
        return f"{formatted_price} USD"
    elif currency == "Rupees":
        return f"Rs. {formatted_price}"
    return formatted_price


def track_changes(inputs):
    """Check if the current inputs differ from the last saved inputs."""
    if st.session_state.last_inputs != inputs:
        st.session_state.download_visible = False
    st.session_state.last_inputs = inputs

def generate_invoice():
    """Streamlit app for generating invoices."""
    st.title("Invoice Generator")

    region = st.selectbox("Region", ["ROW", "India"])
    client_name = st.text_input("Client Name")
    company_name = st.text_input("Company Name")
    contact = st.text_input("Contact Number")
    address = st.text_area("Address")
    project_name = st.text_input("Project Name")
    email = st.text_input("Email")
    service = st.text_input("Service")
    currency = st.selectbox("Currency", ["USD", "Rupees"])
    total_amount = st.number_input("Total Amount", min_value=0.0, format="%.2f")
    payment_option = st.selectbox("Payment Option", ["One Part", "Two Parts", "Three Parts"])
    invoice_date = st.date_input("Invoice Date", value=datetime.today())

    service_description = ""
    if payment_option == "One Part":
        service_description = st.text_area("Service Description (Optional)")

    if payment_option == "Two Parts":
        p1_percentage = st.number_input("Percentage for First Installment", min_value=0.0, max_value=100.0)
        p2_percentage = 100 - p1_percentage
    elif payment_option == "Three Parts":
        p1_percentage = st.number_input("Percentage for First Installment", min_value=0.0, max_value=100.0)
        max_p2 = 100 - p1_percentage
        p2_percentage = st.number_input("Percentage for Second Installment", min_value=0.0, max_value=max_p2)
        p3_percentage = 100 - (p1_percentage + p2_percentage)
        
    if payment_option == "Two Parts":
       p1_percentage = round(p1_percentage)
       p2_percentage = 100 - p1_percentage
       price = round(total_amount * (p1_percentage / 100))
       price2 = total_amount - price
       
    elif payment_option == "Three Parts":
       p1_percentage = round(p1_percentage)
       p2_percentage = round(p2_percentage)
       p3_percentage = 100 - (p1_percentage + p2_percentage)
       price = round(total_amount * (p1_percentage / 100))
       price2 = round(total_amount * (p2_percentage / 100))
       price3 = total_amount - (price + price2)      

    formatted_date = invoice_date.strftime("%d/%m/%Y")
    placeholders = {
        "<< Client Name >>": client_name,
        "<<Company Name>>": company_name,
        "<<Client Contact>>": contact,
        "<<Address>>": address,
        "<<Client Email>>": email,
        "<<Project Name>>": project_name,
        "<<Service>>": service,
        "<<Price>>": format_price(total_amount, currency),
        "<< Date >>": formatted_date,
        "<<Total Amount>>": format_price(total_amount, currency),
    }

    if service_description:
        placeholders["<<Service Description>>"] = service_description

    if payment_option == "Two Parts":
        placeholders.update({
            "<<P1>>": format_percentage(p1_percentage),
            "<<Price>>": format_price(price, currency),
            "<<P2>>": format_percentage(p2_percentage),
            "<<Price2>>": format_price(price2, currency),
       })
    elif payment_option == "Three Parts":
        placeholders.update({
            "<<P1>>": format_percentage(p1_percentage),
            "<<Price>>": format_price(price, currency),
            "<<P2>>": format_percentage(p2_percentage),
            "<<Price2>>": format_price(price2, currency),
            "<<P3>>": format_percentage(p3_percentage),
            "<<Price3>>": format_price(price3, currency),
        })
        
    current_inputs = {
        "region": region,
        "client_name": client_name,
        "company_name": company_name,
        "contact": contact,
        "address": address,
        "project_name": project_name,
        "email": email,
        "service": service,
        "currency": currency,
        "total_amount": total_amount,
        "payment_option": payment_option,
        "invoice_date": invoice_date,
        "service_description": service_description,
    }
#HI
    track_changes(current_inputs)
    

    if st.button("Generate Invoice"):
        invoice_number = get_next_invoice_number()
        placeholders["<<Invoice>>"] = str(invoice_number)
        
        if payment_option == "One Part" and not service_description.strip():
            # Use no-service templates if service description is empty
            template_name = {
                "ROW": "One Part Payment ROW no service.docx",
                "India": "One Part Payment INDIA no service.docx",
            }[region]
        else:
         template_name = {
            "One Part": {
                "ROW": "One Part Payment ROW.docx",
                "India": "One Part Payment INDIA.docx",
            },
            "Two Parts": {
                "ROW": "Two Parts Payment ROW.docx",
                "India": "Two Parts Payment INDIA.docx",
            },
            "Three Parts": {
                "ROW": "Three Parts Payment ROW.docx",
                "India": "Three Parts Payment INDIA.docx",
            },
        }[payment_option][region]

        formatted_date_filename = invoice_date.strftime("%d %b %Y")
        st.session_state.output_path = f"Invoice - {client_name} {formatted_date_filename}.docx"
        try:
            edit_invoice_template(template_name, st.session_state.output_path, placeholders)
            st.session_state.download_visible = True
            st.success(f"Invoice #{invoice_number} generated successfully!")
        except Exception as e:
            st.error(f"An error occurred: {e}")

    if st.session_state.download_visible and st.session_state.output_path:
        with open(st.session_state.output_path, "rb") as file:
            st.download_button(
                label="Download Invoice",
                data=file,
                file_name=st.session_state.output_path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )



# Streamlit App
st.title("Dynamic Document Generator")
option = st.selectbox("Select Document Type", ["NDA", "Contract", "Pricing List","Invoice"], key="doc_type")

base_dir = os.path.abspath(os.path.dirname(__file__))
if option in ["NDA", "Contract"]:
    region = st.selectbox("Region", ["India", "ROW"], key="region")
    template_path = os.path.join(base_dir, f"{option} Template - {'INDIA 3' if region == 'India' else 'ROW 3'}.docx")

    client_name = st.text_input("Enter Client Name:", key="client_name")
    company_name = st.text_input("Enter Company Name:", key="company_name")
    address = st.text_area("Enter Address:", key="address")
    date_field = st.date_input("Enter Date:", datetime.today(), key="date_field")

    placeholders = {
        "<< Client Name >>": client_name,
        "<<Company Name>>": company_name,
        "<<Address>>": address,
        "<< Date (Signature) >>": date_field.strftime("%d-%m-%Y"),
        "<< Date >>": date_field.strftime("%d-%m-%Y"),
    }

elif option == "Pricing List":
    currency = st.selectbox("Select Currency", ["USD", "Rupees", "Pounds"], key="currency")
    if currency == "USD":
        template_path = os.path.join(base_dir, "DM & Automations Services Pricing - USD.docx")
    elif currency == "Rupees":
        template_path = os.path.join(base_dir, "DM & Automations Services Pricing - Rupees.docx")
    elif currency == "Pounds":
        template_path = os.path.join(base_dir, "DM & Automations Services Pricing - Pounds.docx")

    client_name = st.text_input("Enter Client Name:", key="client_name_pricing")
    designation = st.text_input("Enter Designation:", key="designation")
    contact = st.text_input("Enter Contact Number:", key="contact")
    email = st.text_input("Enter Email ID:", key="email")
    location = st.selectbox("Location", ["India", "ROW"], key="location")
    date_field = st.date_input("Enter Date:", datetime.today(), key="date_field_pricing")
    select_all_services = st.checkbox("Select All Services", key="select_all_services")
    services = [
        "Landing page website (design + development)",
        "AI Automations (6 Scenarios)",
        "Whatsapp Automation + Whatsapp Cloud Business Account Setup",
        "CRM Setup",
        "Email Marketing Setup",
        "Make/Zapier Automation",
        "Firefly Meeting Automation",
        "Marketing Strategy",
        "Social Media Channels",
        "Creatives (10 Per Month)",
        "Creatives (20 Per Month)",
        "Creatives (30 Per Month)",
        "Reels (10 Reels)",
        "Meta Ad Account Setup & Pages Setup",
        "Paid Ads (Lead Generation)",
        "Monthly Maintenance & Reporting",
        "AI Chatbot",
        "PDF Generation Automations",
        "AI Generated Social Media Content & Calendar",
        "Custom AI Models & Agents"
    ]
    if select_all_services:
        selected_services = services
    else:
        selected_services = st.multiselect("Select Services", services, key="selected_services")

    placeholders = {
        "<<Client Name>>": client_name,
        "<<Client Designation>>": designation,
        "<<Client Contact>>": contact,
        "<<Client Email>>": email,
        "<<Client Location>>": location,
        "<< Date >>": date_field.strftime("%d-%m-%Y"),
    }
if option == "Invoice":
    generate_invoice()
if option !="Invoice":    
 if st.button("Generate Document", key="generate_button"):
    current_date_str = datetime.now().strftime("%d_%b_%Y").lower()
    
    file_type = {
        "NDA": "NDA Agreement",
        "Contract": "Contract Agreement",
        "Pricing List": "Pricing List"
    }[option]

    formatted_date = date_field.strftime("%d %b %Y")  # e.g., 10 Dec 2024
    file_name = f"{file_type} - {client_name} {formatted_date}.docx"
    word_output_path = os.path.join(base_dir, file_name)

    try:
        if option == "Pricing List":
            updated_word_path = edit_pricing_template(
                template_path, word_output_path, client_name, designation, contact, email, location, selected_services
            )
        else:
            font_size = 11 if option == "NDA" else 12
            updated_word_path = edit_word_template(
                template_path, word_output_path, placeholders, "Times New Roman", font_size, option
            )

        st.success(f"{option} Document Generated Successfully!")

        # Add download button
        with open(updated_word_path, "rb") as word_file:
            st.download_button(
                label="Download Word Document",
                data=word_file,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    except Exception as e:
        st.error(f"An error occurred: {e}")